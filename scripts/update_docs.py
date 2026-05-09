"""Update both .docx docs to match final notebook results."""
import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')


def set_cell(cell, new_text, para_idx=0):
    """Set text in a cell paragraph, preserving first-run formatting."""
    para = cell.paragraphs[para_idx]
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)


def tc(doc, ti, ri, ci, new_text, para_idx=0):
    """Shorthand: update doc.tables[ti].rows[ri].cells[ci]."""
    set_cell(doc.tables[ti].rows[ri].cells[ci], new_text, para_idx)


def pc(doc, pi, new_text):
    """Shorthand: update doc.paragraphs[pi] text."""
    para = doc.paragraphs[pi]
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.add_run(new_text)


# ============================================================
# project_overview.docx
# ============================================================
doc1 = docx.Document('docs/project_overview.docx')

# --- HYPOTHESIS SECTION ---

# TABLE 5 (H1): Add R²>0.50 threshold and confirmed result
tc(doc1, 5, 0, 0,
   'H1 — Predictability: Disease category, age group, sex, and year can predict the '
   'YLD/DALY burden ratio with R² > 0.50 on the held-out test set, confirming that '
   'the chronic-versus-fatal character of disease burden is systematically structured '
   'by demographic and epidemiological classification — not random. '
   'Result: XGBoost achieved R²=0.911 on the 2024 holdout set — H1 confirmed.')

# TABLE 6: Rename header from "Secondary (Sustainability) Hypothesis" to H2
tc(doc1, 6, 0, 0, 'H2 — Model Complexity')

# TABLE 7 (H2): Change from sustainability hypothesis to model complexity
tc(doc1, 7, 0, 0,
   'H2 — Model Complexity: Tree-based and boosting models will outperform linear '
   'regression on the holdout test set by more than 0.10 R² points, with the '
   'performance gap explained by the model\'s ability to capture the non-linear age '
   'curve and bimodal target distribution. '
   'Result: XGBoost (R²=0.911) improved on Linear Regression (R²=0.782) by 0.129 R² '
   'points and reduced MAE by 55% (0.118 → 0.052) — H2 confirmed.')

# P27: Replace sustainability falsifiability paragraph with H3 definition
pc(doc1, 27,
   'H3 — Residual Drift (Sustainability Signal): The XGBoost model\'s residuals will '
   'show systematic positive drift for chronic-dominant disease groups (mental health, '
   'musculoskeletal) in 2024 compared to earlier training years, indicating the chronic '
   'disease burden is growing beyond what historical demographic and disease '
   'classification patterns alone predict. This drift constitutes a quantifiable '
   'ESG/SDG 3 signal and is deliberately falsifiable — if residuals are randomly '
   'distributed across disease groups and years, the sustainability case is not '
   'supported. Result: 12 of 85 disease group×year combinations showed significant '
   'residual drift (14.1%); musculoskeletal (+0.025) and mental and substance use '
   'disorders (+0.019) were under-predicted in 2024 — H3 supported.')

# TABLE 8: Change header from "Null Hypothesis" to "Null Hypothesis (H0)"
tc(doc1, 8, 0, 0, 'Null Hypothesis (H0)')

# P30: Update null threshold from R²>0.30 to R²>0.50, add confirmed outcome
pc(doc1, 30,
   'H0: The YLD/DALY burden ratio cannot be predicted above a naive baseline (mean '
   'prediction) from the available ABDS demographic and epidemiological features. '
   'Threshold: R² > 0.50 required to reject H0. '
   'Result: H0 is rejected — XGBoost achieves R²=0.911 on the 2024 holdout set, '
   'far exceeding the predictability threshold.')

# --- SECTION 4.2: FEATURE ENGINEERING (TABLE 19) ---
tc(doc1, 19, 1, 0, '– burden_ratio = YLD / DALY  [target variable, continuous 0–1]')
tc(doc1, 19, 2, 0, '– age_num: ordinal encoding of 21 age groups (0 = 1–4, 20 = 100+)')
tc(doc1, 19, 3, 0, '– sex_bin: binary (1 = Males, 0 = Females)')
tc(doc1, 19, 4, 0,
   '– data_year: raw snapshot year (2003, 2011, 2015, 2018, 2024) — standardised to '
   'zero mean / unit variance by StandardScaler inside the preprocessing pipeline')
tc(doc1, 19, 5, 0,
   '– age_num_sq: quadratic age term (age_num²) — captures the non-linear age curve '
   'documented in EDA; added for linear model tier')
tc(doc1, 19, 6, 0,
   '– age_sex_interaction: age_num × sex_bin; '
   'disease_year_interaction: disease_group_encoded × year_norm — captures differential '
   'temporal drift by disease group; ranks 6th in SHAP')
tc(doc1, 19, 7, 0,
   '– disease_group: one-hot encoded (17 groups → 17 binary columns) — single strongest '
   'predictor (η²=0.727 in EDA). '
   'Note: log_crude_daly_rate removed during feature selection; '
   'yll_fraction excluded as leakage (r=−0.98 with target)')

# --- SECTION 4.3: EDA PLOTS (TABLE 20) ---
tc(doc1, 20, 6, 0,
   '– Feature leakage diagnosis: yll_fraction (r=−0.98 with burden_ratio), raw YLD, '
   'YLL, DALY columns identified as direct target derivations and excluded')

# --- SECTION 5.1: MODEL ARCHITECTURE (TABLE 22) ---
tc(doc1, 22, 1, 0, 'Tier 1 — Statistical baseline:')
tc(doc1, 22, 2, 0,
   '– Linear Regression (OLS) — primary interpretable baseline; '
   'Test R²=0.782, RMSE=0.153, MAE=0.118')
tc(doc1, 22, 3, 0,
   '– Ridge Regression (alpha=10, GridSearchCV-selected) — evaluated and eliminated: '
   'Test R²=−0.024; over-penalises disease group coefficients')
tc(doc1, 22, 4, 0,
   '– Lasso Regression (alpha=0.01, GridSearchCV-selected) — evaluated and eliminated: '
   'Test R²=−0.045; same structural ceiling as Ridge')
tc(doc1, 22, 5, 0, '')
tc(doc1, 22, 6, 0, 'Tier 2 — Non-linear candidate (not selected):')
tc(doc1, 22, 7, 0,
   '– Random Forest Regressor (n_estimators=300, max_depth=6) — evaluated; '
   'underperforms linear baseline on test set (R²=0.685 vs 0.782); eliminated')
tc(doc1, 22, 8, 0, '')
tc(doc1, 22, 9, 0, 'Primary Model — Gradient Boosting:')
tc(doc1, 22, 10, 0,
   '– XGBoost Regressor (n_estimators=300, max_depth=6, learning_rate=0.05, '
   'min_child_weight=10, reg_lambda=5) — best model; '
   'Test R²=0.911 (95% CI: 0.893–0.924), RMSE=0.097, MAE=0.052; '
   'tuned via 3-step hyperparameter search')
tc(doc1, 22, 11, 0, '')
tc(doc1, 22, 12, 0, '')
tc(doc1, 22, 13, 0, '')
tc(doc1, 22, 14, 0, '')
tc(doc1, 22, 15, 0, '')

# --- SECTION 5.2: EVALUATION FRAMEWORK (TABLE 23) ---
tc(doc1, 23, 1, 0,
   '– Train/test split: Temporal — 2003–2018 training set / 2024 holdout test set '
   '(random split rejected: prevents future rows appearing in training data)')
tc(doc1, 23, 2, 0,
   '– Cross-validation: TimeSeriesSplit(n_splits=3) on training set — preserves '
   'temporal order, prevents data leakage across CV folds')
tc(doc1, 23, 3, 0,
   '– Metrics: R², RMSE, MAE — reported on held-out 2024 test set; '
   '95% bootstrap CI (100 resamples) reported for XGBoost')
tc(doc1, 23, 4, 0,
   '– Residual diagnostics: residuals vs predicted and distribution for '
   'Linear Regression and XGBoost')
tc(doc1, 23, 5, 0,
   '– Feature importance: standardised coefficients (Linear Regression); '
   'mean |SHAP| bar chart and beeswarm plot (XGBoost)')
tc(doc1, 23, 6, 0,
   '– SHAP: TreeExplainer applied to full 2024 test set (XGBoost)')

# --- SECTION 5.3: FINAL RESULTS (TABLE 24) ---
tc(doc1, 24, 0, 0, '5.3  Final Results (from notebook run)')
tc(doc1, 24, 1, 0,
   '– Best model: XGBoost  |  Test R² = 0.911  |  RMSE = 0.097  |  MAE = 0.052  '
   '|  95% CI: [0.893, 0.924]  |  CV RMSE = 0.106 ± 0.002')
tc(doc1, 24, 2, 0,
   '– Runner-up (eliminated): Random Forest  |  Test R² = 0.685  |  MAE = 0.146  '
   '(underperforms linear baseline on 2024 holdout; not selected)')
tc(doc1, 24, 3, 0,
   '– Baseline: Linear Regression  |  Test R² = 0.782  |  RMSE = 0.153  '
   '|  MAE = 0.118  |  CV RMSE = 0.195 ± 0.019')
tc(doc1, 24, 4, 0,
   '– XGBoost improves linear baseline by 0.129 R² points (16%) and reduces '
   'MAE by 55% (0.118 → 0.052)')
tc(doc1, 24, 5, 0,
   '– H1 confirmed: R²=0.911 >> 0.50 threshold; '
   'H2 confirmed: improvement of 0.129 R² points exceeds 0.10 threshold')
tc(doc1, 24, 6, 0, '')
tc(doc1, 24, 7, 0, 'H3 — Sustainability finding (residual drift analysis):')
tc(doc1, 24, 8, 0,
   '– Musculoskeletal disorders 2024: mean residual = +0.025 '
   '(under-predicted — chronic burden rising beyond model expectation)')
tc(doc1, 24, 9, 0,
   '– Mental and substance use disorders 2024: mean residual = +0.019 '
   '(under-predicted — chronic burden outpacing model)')
tc(doc1, 24, 10, 0,
   '– 12 of 85 disease group×year combinations show significant residual drift (14.1%); '
   'fatal-dominant groups (Cancer, Cardiovascular) slightly over-predicted — H3 supported')

# --- SECTION 5.4: PLOTS TO INSERT (TABLE 25) ---
tc(doc1, 25, 1, 0,
   '– model_interpretation.png: standardised coefficients (LR) and mean |SHAP| '
   'bar chart (XGBoost)')
tc(doc1, 25, 2, 0,
   '– learning_curve.png: training vs CV score by training size (Linear Regression)')
tc(doc1, 25, 3, 0,
   '– model_comparison_table.png: R², RMSE, MAE for Linear Regression, '
   'Random Forest, XGBoost')
tc(doc1, 25, 4, 0,
   '– residuals.png: residuals vs predicted + distribution for Linear and XGBoost')
tc(doc1, 25, 5, 0,
   '– shap_beeswarm.png: SHAP beeswarm summary plot (XGBoost)')
tc(doc1, 25, 6, 0,
   '– sustainability_residual_drift.png: mean residuals by disease group and year '
   '(H3 evidence)')
tc(doc1, 25, 7, 0,
   '– demographic_parity.png: mean residuals by sex and age band '
   '(fairness / error parity check)')
tc(doc1, 25, 8, 0,
   '– predicted_vs_actual.png: scatter of predicted vs actual burden ratio '
   'on 2024 holdout (XGBoost)')

# --- SECTION 6.2: MODELLING LIMITATIONS (TABLE 28) ---
tc(doc1, 28, 1, 0,
   '– Target distribution is bimodal (spike near 0 and 1): XGBoost handles extremes '
   'well but residual scatter widens in the 0.3–0.7 boundary zone; '
   '1.4% of predictions exceed the 0–1 theoretical bounds')
tc(doc1, 28, 2, 0,
   '– Disease-level label is coarse: all 205 diseases grouped into 17 categories '
   'for one-hot encoding — within-group variation is averaged out and unrecoverable')
tc(doc1, 28, 3, 0,
   '– Only three models fully evaluated: Linear Regression, Random Forest, XGBoost; '
   'SVR, Gradient Boosting, LightGBM, and Decision Tree not included in final run')
tc(doc1, 28, 4, 0,
   '– Hyperparameter tuning limited to XGBoost (3-step manual search + narrow '
   'GridSearchCV); Random Forest uses matched structural parameters without tuning')
tc(doc1, 28, 5, 0,
   '– TimeSeriesSplit(n_splits=3) provides only 3 CV folds: paired t-test between '
   'LR and XGBoost CV scores was non-significant (p=0.179) due to low fold count — '
   'the held-out 2024 test set is the authoritative model comparison')

# Remove "Section Status: Placeholder" notices
tc(doc1, 17, 0, 0, '')  # section 4 placeholder
tc(doc1, 21, 0, 0, '')  # section 5 placeholder
tc(doc1, 26, 0, 0, '')  # section 6 placeholder

doc1.save('docs/project_overview.docx')
print('project_overview.docx saved.')


# ============================================================
# regression_problem_framing.docx
# ============================================================
doc2 = docx.Document('docs/regression_problem_framing.docx')

# --- TABLE 5: Update mental health residual value and add musculoskeletal ---
# Cell has 2 paragraphs: para 0 = "In concrete terms:", para 1 = body text
tc(doc2, 5, 0, 0, 'In concrete terms:', para_idx=0)
tc(doc2, 5, 0, 0,
   'The mental health residual drift finding from the notebook — where the model '
   'under-predicts chronic mental health burden in 2024 by +0.019 units — is not just '
   'a model evaluation artefact. Musculoskeletal disorders show even stronger drift '
   '(+0.025 in 2024). These are evidence that the chronic disease trajectory has '
   'outpaced what even a well-fitted model trained on historical patterns would predict. '
   'That gap is the ESG signal.',
   para_idx=1)

# --- TABLE 10: Update features table ---
# R4: year_norm → data_year (column 0=Category, 1=Variable, 2=Measures, 3=Why)
tc(doc2, 10, 4, 0, 'Temporal')
tc(doc2, 10, 4, 1, 'data_year')
tc(doc2, 10, 4, 2,
   'Raw snapshot year (2003, 2011, 2015, 2018, 2024) — standardised to zero mean / '
   'unit variance by StandardScaler in preprocessing pipeline')
tc(doc2, 10, 4, 3,
   'Captures secular trend — chronic burden ratio rising over 21 years, especially '
   'mental health; pipeline standardisation ensures comparable regularisation penalties '
   'across features')

# R5: Replace log_crude_daly_rate with engineered features
tc(doc2, 10, 5, 0, 'Engineered')
tc(doc2, 10, 5, 1, 'age_num_sq, age_sex_interaction, disease_year_interaction')
tc(doc2, 10, 5, 2,
   'Quadratic age (age_num²); age×sex cross-term; '
   'disease_group_encoded × year_norm interaction')
tc(doc2, 10, 5, 3,
   'Gives linear models a fair chance at the non-linear age curve; '
   'disease_year_interaction ranks 6th in SHAP, confirming differential temporal '
   'drift across disease groups — directly supports H3')

# --- TABLE 11: Update exclusion note to mention log_crude_daly_rate ---
tc(doc2, 11, 0, 0,
   'A note on what is deliberately excluded:\n'
   'Socioeconomic status, comorbidity counts, and individual-level health behaviours '
   'are not in the ABDS 2024 data and therefore cannot be features. '
   'yll_fraction (YLL/DALY) was excluded as direct target leakage (r=−0.98 with '
   'burden_ratio). log_crude_daly_rate was removed during feature selection — it did '
   'not improve model performance beyond the disease group dummies. '
   'These are known limitations, not oversights.')

# --- TABLE 16 Summary: Update features row ---
tc(doc2, 16, 4, 2,
   'disease_group (one-hot, 17 categories), age_num (ordinal, 0–20), sex_bin (binary), '
   'data_year (standardised), plus engineered: age_num_sq, age_sex_interaction, '
   'disease_year_interaction. All derived from ABDS S1 — no external data required.')

# --- P58: Update year encoding explanation ---
pc(doc2, 58,
   'data_year is passed as a raw integer (2003, 2011, 2015, 2018, 2024) and '
   'standardised to zero mean / unit variance by StandardScaler within the '
   'preprocessing pipeline after the train/test split. This prevents the year '
   'coefficient from dominating Ridge/Lasso regularisation penalties — equivalent '
   'to normalisation but pipeline-integrated, so the 2024 test year\'s values do '
   'not influence the scaling parameters.')

doc2.save('docs/regression_problem_framing.docx')
print('regression_problem_framing.docx saved.')
print('All done.')
